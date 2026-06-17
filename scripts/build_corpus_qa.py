"""Build a text/chunk corpus and QA report from downloaded student resources.

The script walks:

- data/blackboard
- data/rencai/raw
- data/transcripts/raw

It writes normalized extracted text, chunk JSONL, and file-level QA reports under
data/corpus. Generated corpus outputs are intended to stay local unless reviewed.
"""

from __future__ import annotations

import argparse
import csv
import hashlib
import json
import logging
import re
import shutil
import zipfile
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path
from typing import Any, Callable


TEXT_SUFFIXES = {".txt", ".md", ".csv", ".json", ".rtf"}
TRANSCRIPT_SUFFIXES = {".srt", ".vtt"}
IMAGE_SUFFIXES = {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"}
SKIP_NAMES = {".gitkeep", "README.md"}
MOJIBAKE_MARKERS = (
    "Ã",
    "Â",
    "â€",
    "ï¼",
    "ï",
    "ã€",
    "å®",
    "å¤",
    "å›",
    "ä¸",
    "äº",
    "æœ",
    "ç”",
    "è¯",
    "é‡",
)
STOPWORDS = {
    "about",
    "after",
    "again",
    "also",
    "and",
    "are",
    "because",
    "been",
    "before",
    "being",
    "between",
    "can",
    "china",
    "for",
    "from",
    "has",
    "have",
    "how",
    "into",
    "its",
    "may",
    "not",
    "our",
    "scholar",
    "scholars",
    "schwarzman",
    "students",
    "student",
    "that",
    "the",
    "their",
    "this",
    "through",
    "will",
    "with",
    "you",
    "your",
}


@dataclass
class ExtractResult:
    status: str
    text: str = ""
    error: str = ""
    detected_type: str = ""
    notes: list[str] = field(default_factory=list)
    metadata: dict[str, Any] = field(default_factory=dict)


@dataclass
class SourceFile:
    source: str
    path: Path
    display_path: str
    source_path: str


def sha256_file(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for block in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(block)
    return digest.hexdigest()


def safe_part(value: str, fallback: str = "untitled") -> str:
    cleaned = re.sub(r"[^A-Za-z0-9_. -]+", "_", value).strip(" .")
    cleaned = re.sub(r"\s+", " ", cleaned)
    return (cleaned or fallback)[:110]


def rel_posix(path: Path, root: Path) -> str:
    return str(path.relative_to(root)).replace("\\", "/")


def clean_text(text: str) -> str:
    text = repair_mojibake(text)
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    lines = [line.strip() for line in text.splitlines()]
    return "\n".join(line for line in lines if line).strip()


def detect_type(path: Path) -> str:
    suffix = path.suffix.lower()
    try:
        head = path.read_bytes()[:4096]
    except Exception:
        head = b""

    if head.startswith(b"%PDF"):
        return "pdf"
    if head.startswith(b"\x89PNG\r\n\x1a\n"):
        return "image"
    if head.startswith(b"\xff\xd8\xff"):
        return "image"
    if head.startswith(b"\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1"):
        return "legacy_office"
    if head.lstrip().lower().startswith((b"<!doctype html", b"<html")):
        return "html"
    if zipfile.is_zipfile(path):
        try:
            with zipfile.ZipFile(path) as archive:
                names = set(archive.namelist())
            if "word/document.xml" in names:
                return "docx"
            if "xl/workbook.xml" in names:
                return "xlsx"
            return "zip"
        except Exception:
            return "zip"
    if suffix == ".pdf":
        return "pdf"
    if suffix == ".docx":
        return "docx"
    if suffix == ".xlsx":
        return "xlsx"
    if suffix == ".doc":
        return "legacy_office"
    if suffix in IMAGE_SUFFIXES:
        return "image"
    if suffix in TEXT_SUFFIXES:
        return "text"
    if suffix in TRANSCRIPT_SUFFIXES:
        return "transcript"
    if suffix == ".bin":
        return "binary"
    return suffix.lstrip(".") or "unknown"


def extract_pdf(path: Path) -> ExtractResult:
    try:
        from pypdf import PdfReader  # type: ignore
    except Exception as exc:
        return ExtractResult("missing_dependency", error=f"pypdf not installed: {exc}", detected_type="pdf")

    try:
        reader = PdfReader(str(path))
        parts = []
        for page in reader.pages:
            parts.append(page.extract_text() or "")
        text = clean_text("\n\n".join(parts))
        return ExtractResult("ok", text=text, detected_type="pdf", metadata={"pages": len(reader.pages)})
    except Exception as exc:
        return ExtractResult("error", error=str(exc), detected_type="pdf")


def extract_docx(path: Path) -> ExtractResult:
    try:
        import docx  # type: ignore
    except Exception as exc:
        return ExtractResult("missing_dependency", error=f"python-docx not installed: {exc}", detected_type="docx")

    try:
        document = docx.Document(str(path))
        parts = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return ExtractResult("ok", text=clean_text("\n".join(parts)), detected_type="docx")
    except Exception as exc:
        return ExtractResult("error", error=str(exc), detected_type="docx")


def extract_xlsx(path: Path) -> ExtractResult:
    try:
        from openpyxl import load_workbook  # type: ignore
    except Exception as exc:
        return ExtractResult("missing_dependency", error=f"openpyxl not installed: {exc}", detected_type="xlsx")

    try:
        workbook = load_workbook(str(path), read_only=True, data_only=True)
        lines = []
        for sheet in workbook.worksheets:
            lines.append(f"# {sheet.title}")
            for row in sheet.iter_rows(values_only=True):
                values = [str(value).strip() for value in row if value not in (None, "")]
                if values:
                    lines.append("\t".join(values))
        return ExtractResult(
            "ok",
            text=clean_text("\n".join(lines)),
            detected_type="xlsx",
            metadata={"sheets": [sheet.title for sheet in workbook.worksheets]},
        )
    except Exception as exc:
        return ExtractResult("error", error=str(exc), detected_type="xlsx")


def extract_image(path: Path, use_ocr: bool) -> ExtractResult:
    if not use_ocr:
        return ExtractResult("unsupported", error="Image OCR disabled", detected_type="image")
    if not shutil.which("tesseract"):
        return ExtractResult("unsupported", error="Tesseract OCR binary not found", detected_type="image")
    try:
        from PIL import Image  # type: ignore
        import pytesseract  # type: ignore
    except Exception as exc:
        return ExtractResult("missing_dependency", error=f"Pillow/pytesseract not installed: {exc}", detected_type="image")

    try:
        with Image.open(path) as image:
            text = pytesseract.image_to_string(image)
        return ExtractResult("ok", text=clean_text(text), detected_type="image", notes=["ocr"])
    except Exception as exc:
        return ExtractResult("error", error=str(exc), detected_type="image")


def extract_text_file(path: Path) -> ExtractResult:
    try:
        return ExtractResult("ok", text=clean_text(path.read_text(encoding="utf-8", errors="ignore")), detected_type="text")
    except Exception as exc:
        return ExtractResult("error", error=str(exc), detected_type="text")


def extract_transcript_file(path: Path) -> ExtractResult:
    result = extract_text_file(path)
    result.detected_type = "transcript"
    if result.status != "ok":
        return result

    lines: list[str] = []
    previous = ""
    for raw_line in result.text.splitlines():
        line = raw_line.strip().lstrip("\ufeff")
        if not line:
            continue
        if line.upper().startswith("WEBVTT"):
            continue
        if re.fullmatch(r"\d+", line):
            continue
        if "-->" in line and re.search(r"\d{1,2}:\d{2}", line):
            continue
        if re.fullmatch(r"NOTE\b.*", line, flags=re.I):
            continue
        line = re.sub(r"<[^>]+>", "", line).strip()
        line = re.sub(r"\s+", " ", line)
        if not line or line == previous:
            continue
        lines.append(line)
        previous = line

    result.text = clean_text("\n".join(lines))
    result.notes.append("transcript_cleaned")
    return result


def extract_legacy_office(path: Path) -> ExtractResult:
    return ExtractResult(
        "unsupported",
        error="Legacy .doc/.xls file; convert to .docx/.xlsx for reliable extraction",
        detected_type="legacy_office",
    )


def extract_html(path: Path) -> ExtractResult:
    result = extract_text_file(path)
    result.detected_type = "html"
    result.status = "review"
    result.notes.append("downloaded_html")
    result.error = "File appears to be HTML, not a source document"
    return result


def extractor_for(detected_type: str, use_ocr: bool) -> Callable[[Path], ExtractResult]:
    if detected_type == "pdf":
        return extract_pdf
    if detected_type == "docx":
        return extract_docx
    if detected_type == "xlsx":
        return extract_xlsx
    if detected_type == "image":
        return lambda path: extract_image(path, use_ocr)
    if detected_type == "text":
        return extract_text_file
    if detected_type == "transcript":
        return extract_transcript_file
    if detected_type == "html":
        return extract_html
    if detected_type == "legacy_office":
        return extract_legacy_office
    return lambda _: ExtractResult("unsupported", error=f"Unsupported detected type: {detected_type}", detected_type=detected_type)


def iter_source_files(root: Path) -> list[SourceFile]:
    sources: list[SourceFile] = []
    blackboard = root / "data" / "blackboard"
    if blackboard.exists():
        for path in sorted(blackboard.rglob("*")):
            if path.is_file() and path.name not in SKIP_NAMES:
                sources.append(
                    SourceFile(
                        source="blackboard",
                        path=path,
                        display_path=rel_posix(path, root),
                        source_path=rel_posix(path, blackboard),
                    )
                )

    rencai_raw = root / "data" / "rencai" / "raw"
    if rencai_raw.exists():
        for path in sorted(rencai_raw.rglob("*")):
            if path.is_file() and path.name not in SKIP_NAMES:
                sources.append(
                    SourceFile(
                        source="rencai",
                        path=path,
                        display_path=rel_posix(path, root),
                        source_path=rel_posix(path, rencai_raw),
                    )
                )

    transcripts_raw = root / "data" / "transcripts" / "raw"
    if transcripts_raw.exists():
        for path in sorted(transcripts_raw.rglob("*")):
            if path.is_file() and path.name not in SKIP_NAMES:
                sources.append(
                    SourceFile(
                        source="transcripts",
                        path=path,
                        display_path=rel_posix(path, root),
                        source_path=rel_posix(path, transcripts_raw),
                    )
                )
    return sources


def text_output_path(text_dir: Path, source_file: SourceFile, digest: str) -> Path:
    stem = safe_part(source_file.path.stem, "resource")
    parents = [safe_part(part) for part in Path(source_file.source_path).parts[:-1]]
    return text_dir / source_file.source / Path(*parents) / f"{stem}--{digest[:10]}.txt"


def public_citation_ref(source_file: SourceFile) -> str:
    return f"{source_file.source}/{source_file.source_path}".replace("\\", "/")


def one_sentence_summary(title: str, text: str, source_path: str) -> str:
    title_text = title.replace("_", " ").replace("-", " ")
    title_text = re.sub(r"\s+", " ", title_text).strip(" .")
    sentences = re.split(r"(?<=[.!?])\s+", re.sub(r"\s+", " ", text))
    candidates = [
        sentence.strip()
        for sentence in sentences
        if 45 <= len(sentence.strip()) <= 260 and not looks_like_boilerplate(sentence) and not looks_garbled(sentence)
    ]
    if candidates:
        sentence = candidates[0]
        return sentence if sentence.endswith((".", "!", "?")) else f"{sentence}."

    keyword_source = f"{title_text} {source_path}" if looks_garbled(text[:5000]) else f"{title_text} {source_path} {text}"
    keywords = top_keywords(keyword_source, limit=5)
    if keywords:
        return f"Student-facing material about {title_text}, with recurring terms including {', '.join(keywords)}."
    return f"Student-facing material titled {title_text}."


def looks_like_boilerplate(sentence: str) -> bool:
    lowered = sentence.lower()
    boilerplate = [
        "all rights reserved",
        "copyright",
        "page ",
        "http://",
        "https://",
        "table of contents",
        "agenda",
        "powered by",
    ]
    return any(term in lowered for term in boilerplate)


def looks_garbled(text: str) -> bool:
    if not text:
        return False
    return mojibake_score(text) >= 8


def mojibake_score(text: str) -> int:
    sample = text[:5000]
    marker_hits = sum(sample.count(marker) for marker in MOJIBAKE_MARKERS)
    return marker_hits + (25 if "\ufffd" in sample else 0)


def repair_mojibake(text: str) -> str:
    if mojibake_score(text) < 3:
        return text

    original_score = mojibake_score(text)
    best = text
    best_score = original_score
    for encoding in ("cp1252", "latin1"):
        try:
            candidate = text.encode(encoding).decode("utf-8")
        except UnicodeError:
            continue
        score = mojibake_score(candidate)
        if score < best_score and len(candidate) >= len(text) * 0.75:
            best = candidate
            best_score = score
    return best


def top_keywords(text: str, limit: int = 5) -> list[str]:
    words = re.findall(r"[A-Za-z][A-Za-z]{2,}", text.lower())
    counts = Counter(word for word in words if word not in STOPWORDS)
    return [word for word, _ in counts.most_common(limit)]


def chunk_text(text: str, chunk_size: int, overlap: int) -> list[dict[str, Any]]:
    if not text:
        return []
    chunks = []
    start = 0
    text_len = len(text)
    while start < text_len:
        end = min(start + chunk_size, text_len)
        if end < text_len:
            boundary = max(text.rfind("\n", start, end), text.rfind(". ", start, end), text.rfind(" ", start, end))
            if boundary > start + int(chunk_size * 0.55):
                end = boundary + 1
        chunks.append({"char_start": start, "char_end": end, "text": text[start:end].strip()})
        if end >= text_len:
            break
        start = max(end - overlap, start + 1)
    return [chunk for chunk in chunks if chunk["text"]]


def qa_flags(source_file: SourceFile, result: ExtractResult, text_chars: int, size_bytes: int, duplicate_of: str) -> list[str]:
    flags = []
    suffix = source_file.path.suffix.lower()
    if duplicate_of:
        flags.append("duplicate_content")
    if result.status in {"error", "missing_dependency", "unsupported"}:
        flags.append("not_ready")
    if result.detected_type == "legacy_office":
        flags.append("convert_legacy_office")
    if result.detected_type == "image" and result.status != "ok":
        flags.append("image_needs_ocr_or_manual_review")
    if result.detected_type == "html":
        flags.append("possible_download_or_login_page")
    if suffix == ".bin" and result.detected_type in {"pdf", "docx", "xlsx"}:
        flags.append("extension_mismatch_bin")
    if size_bytes == 0:
        flags.append("empty_file")
    if result.status == "ok" and text_chars == 0:
        flags.append("no_text_extracted")
    elif result.status == "ok" and result.detected_type == "transcript" and text_chars < 200:
        flags.append("short_transcript_review")
    elif result.status == "ok" and text_chars < 400 and result.detected_type in {"pdf", "docx", "xlsx"}:
        flags.append("low_text_review")
    if result.status == "ok" and looks_garbled(result.text):
        flags.append("possible_encoding_garbled")
    if result.detected_type == "pdf" and size_bytes > 500_000 and text_chars < 1000:
        flags.append("possible_scanned_pdf")
    return flags or ["ready"]


def write_csv(path: Path, rows: list[dict[str, Any]]) -> None:
    if not rows:
        path.write_text("", encoding="utf-8")
        return
    with path.open("w", encoding="utf-8", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(rows[0].keys()))
        writer.writeheader()
        writer.writerows(rows)


def write_markdown_summary(path: Path, rows: list[dict[str, Any]], chunk_count: int) -> None:
    status_counts = Counter(row["status"] for row in rows)
    flag_counts = Counter(flag for row in rows for flag in str(row["qa_flags"]).split(";") if flag)
    type_counts = Counter(row["detected_type"] for row in rows)
    source_counts = Counter(row["source"] for row in rows)
    review_rows = [row for row in rows if row["qa_flags"] != "ready"]

    lines = [
        "# Corpus QA Summary",
        "",
        f"- Files scanned: {len(rows)}",
        f"- Chunks written: {chunk_count}",
        f"- Sources: {', '.join(f'{source}={count}' for source, count in sorted(source_counts.items()))}",
        f"- Statuses: {', '.join(f'{status}={count}' for status, count in sorted(status_counts.items()))}",
        f"- Detected types: {', '.join(f'{kind}={count}' for kind, count in sorted(type_counts.items()))}",
        "",
        "## Flags",
        "",
    ]
    lines.extend(f"- {flag}: {count}" for flag, count in sorted(flag_counts.items()))
    lines.extend(["", "## Files Needing Review", ""])
    if review_rows:
        for row in review_rows:
            lines.append(f"- `{row['path']}` - {row['qa_flags']} - {row['error']}")
    else:
        lines.append("- None.")
    lines.extend(["", "## One-Line File Summaries", ""])
    for row in rows:
        lines.append(f"- `{row['path']}`: {row['summary']}")
    path.write_text("\n".join(lines) + "\n", encoding="utf-8")


def main() -> int:
    logging.getLogger("pypdf").setLevel(logging.ERROR)

    parser = argparse.ArgumentParser()
    parser.add_argument("--root", default=".", help="Repository root")
    parser.add_argument("--chunk-size", type=int, default=3000, help="Approximate chunk size in characters")
    parser.add_argument("--overlap", type=int, default=350, help="Chunk overlap in characters")
    parser.add_argument("--ocr", action="store_true", help="Attempt OCR for image files if Tesseract is installed")
    args = parser.parse_args()

    root = Path(args.root).resolve()
    stamp = datetime.now().strftime("%Y-%m-%dT%H-%M-%S")
    corpus_dir = root / "data" / "corpus"
    text_dir = corpus_dir / "text"
    chunks_dir = corpus_dir / "chunks"
    reports_dir = corpus_dir / "reports"
    for directory in (text_dir, chunks_dir, reports_dir):
        directory.mkdir(parents=True, exist_ok=True)

    source_files = iter_source_files(root)
    digest_to_paths: dict[str, list[SourceFile]] = defaultdict(list)
    digests: dict[Path, str] = {}
    for source_file in source_files:
        digest = sha256_file(source_file.path)
        digests[source_file.path] = digest
        digest_to_paths[digest].append(source_file)

    rows: list[dict[str, Any]] = []
    chunks: list[dict[str, Any]] = []

    for source_file in source_files:
        digest = digests[source_file.path]
        duplicates = digest_to_paths[digest]
        duplicate_of = ""
        if len(duplicates) > 1 and duplicates[0].path != source_file.path:
            duplicate_of = duplicates[0].display_path

        detected_type = detect_type(source_file.path)
        result = extractor_for(detected_type, use_ocr=args.ocr)(source_file.path)
        if not result.detected_type:
            result.detected_type = detected_type
        text = result.text
        text_path = ""
        file_chunks = []
        if text:
            output_path = text_output_path(text_dir, source_file, digest)
            output_path.parent.mkdir(parents=True, exist_ok=True)
            output_path.write_text(text, encoding="utf-8")
            text_path = rel_posix(output_path, root)
            file_chunks = chunk_text(text, args.chunk_size, args.overlap)
            citation_ref = public_citation_ref(source_file)
            for index, chunk in enumerate(file_chunks):
                chunks.append(
                    {
                        "chunk_id": f"{digest[:12]}-{index:04d}",
                        "source": source_file.source,
                        "path": source_file.display_path,
                        "source_file": source_file.display_path,
                        "source_title": source_file.path.name,
                        "text_path": text_path,
                        "citation_ref": citation_ref,
                        "chunk_index": index,
                        "char_start": chunk["char_start"],
                        "char_end": chunk["char_end"],
                        "sha256": digest,
                        "text": chunk["text"],
                    }
                )

        size_bytes = source_file.path.stat().st_size
        flags = qa_flags(source_file, result, len(text), size_bytes, duplicate_of)
        rows.append(
            {
                "source": source_file.source,
                "path": source_file.display_path,
                "source_path": source_file.source_path,
                "citation_path": public_citation_ref(source_file),
                "filename": source_file.path.name,
                "extension": source_file.path.suffix.lower(),
                "detected_type": result.detected_type,
                "status": result.status,
                "qa_flags": ";".join(flags),
                "summary": one_sentence_summary(source_file.path.stem, text, source_file.source_path),
                "text_path": text_path,
                "chunks": len(file_chunks),
                "chars": len(text),
                "size_bytes": size_bytes,
                "sha256": digest,
                "duplicate_of": duplicate_of,
                "notes": ";".join(result.notes),
                "metadata": json.dumps(result.metadata, ensure_ascii=True, sort_keys=True),
                "error": result.error,
            }
        )

    chunks_jsonl = chunks_dir / f"corpus-chunks-{stamp}.jsonl"
    with chunks_jsonl.open("w", encoding="utf-8") as handle:
        for chunk in chunks:
            handle.write(json.dumps(chunk, ensure_ascii=False) + "\n")

    report_json = reports_dir / f"corpus-file-report-{stamp}.json"
    report_csv = reports_dir / f"corpus-file-report-{stamp}.csv"
    summary_md = reports_dir / f"corpus-qa-summary-{stamp}.md"
    report_json.write_text(json.dumps(rows, ensure_ascii=False, indent=2), encoding="utf-8")
    write_csv(report_csv, rows)
    write_markdown_summary(summary_md, rows, len(chunks))

    print(f"Scanned {len(rows)} files.")
    print(f"Extracted text from {sum(1 for row in rows if row['text_path'])} files.")
    print(f"Wrote {len(chunks)} chunks.")
    print(f"Wrote {report_csv}")
    print(f"Wrote {summary_md}")
    print(f"Wrote {chunks_jsonl}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
